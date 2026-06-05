# -*- coding: utf-8 -*-
"""
报告生成模块 - Word (.docx) + PDF (带数据可视化图表)
"""

import os
import json
import io
from datetime import datetime


# ========== 评分维度映射 ==========
DIMENSION_LABELS = {
    "hook_strength": "钩子强度",
    "visual_quality": "画面质量",
    "copy_quality": "文案质量",
    "audio_quality": "音频质量",
    "rhythm": "节奏感",
    "conversion_guidance": "转化引导",
    "overall": "综合评分",
}


def load_result(video_dir):
    """从 video_dir 加载 result.json"""
    result_path = os.path.join(video_dir, "result.json")
    if not os.path.exists(result_path):
        return None
    with open(result_path, "r", encoding="utf-8") as f:
        return json.load(f)


def get_scoring_data(result):
    """提取评分数据为 {label: score} 格式"""
    scoring = {}
    ca = result.get("content_analysis", {})
    if not ca.get("analyzed"):
        return scoring
    scores = ca.get("scoring", {})
    for key, label in DIMENSION_LABELS.items():
        if key in scores and isinstance(scores[key], dict):
            scoring[label] = {
                "score": scores[key].get("score", 0),
                "max": scores[key].get("max", 10),
                "comment": scores[key].get("comment", ""),
            }
    return scoring


# =====================================================
# Word 报告
# =====================================================
def generate_word_report(result, output_path):
    """
    生成可编辑的 Word 分析报告
    result: 从 result.json 加载的完整数据
    output_path: 输出 .docx 文件路径
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # ---------- 样式调整 ----------
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ---------- 标题 ----------
    video_name = result.get("video_name", "未知视频")
    title = doc.add_heading("信息流素材分析报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 基本信息
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_para.add_run(f"视频: {video_name}  |  生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # 空行

    # ---------- 一、视频基本信息 ----------
    doc.add_heading("一、视频基本信息", level=1)
    vi = result.get("video_info", {})
    info_data = [
        ("文件名", video_name),
        ("分辨率", f"{vi.get('width', '?')} x {vi.get('height', '?')}"),
        ("时长", f"{vi.get('duration', 0)} 秒"),
        ("帧率", f"{vi.get('fps', 0)} fps"),
        ("文件大小", f"{vi.get('file_size_mb', 0)} MB"),
        ("音频", "有" if vi.get("has_audio") else "无"),
        ("创建时间", result.get("created_at", "")),
        ("分析时间", result.get("analyzed_at", "")),
    ]
    table = doc.add_table(rows=len(info_data), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(info_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
        # 加粗标签列
        for paragraph in table.rows[i].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # ---------- 二、语音转文字 ----------
    doc.add_heading("二、语音转文字结果", level=1)
    transcript = result.get("transcript", {})
    if transcript.get("transcribed"):
        doc.add_paragraph(f"识别模型: {transcript.get('model', '未知')}")
        doc.add_paragraph(f"识别文本:")
        # 原文
        p = doc.add_paragraph()
        run = p.add_run(transcript.get("full_text", ""))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x33, 0x55, 0x99)
        # 分段信息
        segments = transcript.get("segments", [])
        if segments:
            doc.add_paragraph(f"分段数量: {len(segments)}")
            for seg in segments[:10]:  # 最多显示10段
                doc.add_paragraph(
                    f"[{seg.get('start', 0):.1f}s - {seg.get('end', 0):.1f}s] {seg.get('text', '')}",
                    style="List Bullet"
                )
    else:
        doc.add_paragraph("无语音内容或语音识别未执行", style="List Bullet")

    doc.add_paragraph()

    # ---------- 三、画面分析 ----------
    doc.add_heading("三、关键帧画面分析", level=1)
    fa = result.get("frame_analysis", {})
    if fa.get("analyzed"):
        doc.add_paragraph(f"分析模型: {fa.get('model', '未知')}")
        doc.add_paragraph(f"分析帧数: {fa.get('frames_analyzed', 0)} / {fa.get('total_frames', 0)}")

        # 整体视觉总结
        if fa.get("overall_visual"):
            doc.add_heading("整体视觉总结", level=2)
            doc.add_paragraph(fa["overall_visual"])

        # 逐帧分析
        frames = fa.get("frames", [])
        if frames:
            doc.add_heading("逐帧分析", level=2)
            for f in frames:
                p = doc.add_paragraph()
                run = p.add_run(f"第 {f.get('frame', '?')} 帧")
                run.font.bold = True
                doc.add_paragraph(f"内容: {f.get('content', '')}", style="List Bullet")
                doc.add_paragraph(f"文字: {f.get('text', '无')}", style="List Bullet")
                doc.add_paragraph(f"风格: {f.get('style', '')}", style="List Bullet")
                doc.add_paragraph(f"变化: {f.get('change', '')}", style="List Bullet")
    else:
        doc.add_paragraph("画面分析未执行", style="List Bullet")

    doc.add_paragraph()

    # ---------- 四、综合 AI 分析 ----------
    ca = result.get("content_analysis", {})
    if ca.get("analyzed"):
        # 4.1 元素拆解
        doc.add_heading("四、元素拆解", level=1)
        eb = ca.get("element_breakdown", {})

        # 画面
        if eb.get("visual"):
            doc.add_heading("4.1 画面元素", level=2)
            vis = eb["visual"]
            for key, label in [("subject", "画面主体"), ("scene", "场景"),
                               ("color_tone", "色调"), ("composition", "构图"), ("effects", "特效")]:
                if vis.get(key):
                    doc.add_paragraph(f"{label}: {vis[key]}", style="List Bullet")

        # 文案
        if eb.get("text"):
            doc.add_heading("4.2 文案元素", level=2)
            txt = eb["text"]
            for key, label in [("spoken", "口播/旁白"), ("subtitles", "字幕/贴纸"),
                               ("cta", "转化引导"), ("headline", "标题/大字")]:
                if txt.get(key):
                    doc.add_paragraph(f"{label}: {txt[key]}", style="List Bullet")

        # 音频
        if eb.get("audio"):
            doc.add_heading("4.3 音频元素", level=2)
            aud = eb["audio"]
            for key, label in [("voice_style", "语音风格"), ("bgm_style", "BGM风格"),
                               ("sound_effects", "音效")]:
                if aud.get(key):
                    doc.add_paragraph(f"{label}: {aud[key]}", style="List Bullet")

        # 结构
        if eb.get("structure"):
            doc.add_heading("4.4 结构分析", level=2)
            st = eb["structure"]
            for key, label in [("hook", "钩子(前3秒)"), ("body", "中间内容"),
                               ("ending", "结尾/CTA"), ("rhythm", "节奏")]:
                if st.get(key):
                    doc.add_paragraph(f"{label}: {st[key]}", style="List Bullet")

        doc.add_paragraph()

        # 4.2 评分
        doc.add_heading("五、评分详情", level=1)
        scoring = ca.get("scoring", {})
        score_table = doc.add_table(rows=1, cols=4)
        score_table.style = "Light Grid Accent 1"
        score_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = score_table.rows[0].cells
        hdr[0].text = "维度"
        hdr[1].text = "分数"
        hdr[2].text = "满分"
        hdr[3].text = "评语"
        for cell in hdr:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for key in ["hook_strength", "visual_quality", "copy_quality",
                     "audio_quality", "rhythm", "conversion_guidance", "overall"]:
            dim = scoring.get(key, {})
            if isinstance(dim, dict):
                row = score_table.add_row().cells
                row[0].text = DIMENSION_LABELS.get(key, key)
                row[1].text = str(dim.get("score", 0))
                row[2].text = str(dim.get("max", 10))
                row[3].text = dim.get("comment", "")

        doc.add_paragraph()

        # 4.3 跑量归因
        doc.add_heading("六、跑量归因分析", level=1)
        ta = ca.get("traffic_attribution", {})

        # CTR 驱动
        ctr = ta.get("ctr_drivers", [])
        if ctr:
            doc.add_heading("点击率(CTR)驱动因素", level=2)
            for item in ctr:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_heading("点击率(CTR)驱动因素", level=2)
            doc.add_paragraph("无明显CTR驱动因素", style="List Bullet")

        # CVR 驱动
        cvr = ta.get("cvr_drivers", [])
        if cvr:
            doc.add_heading("转化率(CVR)驱动因素", level=2)
            for item in cvr:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_heading("转化率(CVR)驱动因素", level=2)
            doc.add_paragraph("无明显CVR驱动因素", style="List Bullet")

        # 完播率驱动
        comp = ta.get("completion_drivers", [])
        if comp:
            doc.add_heading("完播率驱动因素", level=2)
            for item in comp:
                doc.add_paragraph(item, style="List Bullet")
        else:
            doc.add_heading("完播率驱动因素", level=2)
            doc.add_paragraph("无明显完播率驱动因素", style="List Bullet")

        # 潜在问题
        issues = ta.get("potential_issues", [])
        if issues:
            doc.add_heading("潜在问题", level=2)
            for item in issues:
                doc.add_paragraph(item, style="List Bullet")

        # 综合评估
        if ta.get("overall_assessment"):
            doc.add_heading("综合跑量评估", level=2)
            doc.add_paragraph(ta["overall_assessment"])

        # 改进建议
        suggestions = ta.get("improvement_suggestions", [])
        if suggestions:
            doc.add_heading("改进建议", level=2)
            for i, item in enumerate(suggestions, 1):
                doc.add_paragraph(f"{i}. {item}")
    else:
        doc.add_heading("四、综合 AI 分析", level=1)
        doc.add_paragraph("AI 综合分析尚未执行，请先完成分析后再导出报告。")

    # ---------- 页脚 ----------
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run("— 由信息流素材分析 Agent 自动生成 —")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)

    # 保存
    doc.save(output_path)
    return output_path


# =====================================================
# PDF 报告（带数据可视化）
# =====================================================
def _get_chinese_font_prop():
    """获取中文字体 FontProperties"""
    import matplotlib.font_manager as fm
    # 直接用微软雅黑
    msyh_path = r"C:\Windows\Fonts\msyh.ttc"
    if os.path.exists(msyh_path):
        return fm.FontProperties(fname=msyh_path)
    # 备选
    for fp in fm.findSystemFonts():
        try:
            fn = os.path.basename(fp).lower()
            if "msyh" in fn or "simhei" in fn:
                return fm.FontProperties(fname=fp)
        except Exception:
            pass
    return None


def _get_pdf_font_paths():
    """获取 PDF 可用字体路径，兼容 Windows 和 Linux/Render。"""
    candidates_regular = [
        r"C:\Windows\Fonts\msyh.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    candidates_bold = [
        r"C:\Windows\Fonts\msyhbd.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Bold.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ]

    regular = next((p for p in candidates_regular if os.path.exists(p)), None)
    bold = next((p for p in candidates_bold if os.path.exists(p)), None) or regular
    return regular, bold


def _generate_score_bar_chart(scoring_dict, output_path):
    """
    生成评分柱状图
    scoring_dict: {label: {score, max, comment}}
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    font_prop = _get_chinese_font_prop()

    # 只画6个子维度（不含overall）
    dims = [k for k in scoring_dict if k != "综合评分"]
    scores = [scoring_dict[k]["score"] for k in dims]

    fig, ax = plt.subplots(figsize=(8, 4))
    bars = ax.barh(dims, scores, color=["#7C6AFF" if s >= 7 else "#F39C12" if s >= 4 else "#E74C3C" for s in scores],
                   height=0.6, edgecolor="none")
    ax.set_xlim(0, 10)
    ax.set_xlabel("分数", fontsize=10, fontproperties=font_prop)
    ax.set_title("六维评分", fontsize=14, fontweight="bold", pad=12, fontproperties=font_prop)
    ax.tick_params(axis="y", labelsize=11)
    if font_prop:
        for label in ax.get_yticklabels():
            label.set_fontproperties(font_prop)

    # 在柱状图右侧显示分数
    for bar, score in zip(bars, scores):
        ax.text(bar.get_width() + 0.2, bar.get_y() + bar.get_height() / 2,
                f"{score}", va="center", fontsize=11, fontweight="bold")

    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", alpha=0.3)
    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return output_path


def _generate_radar_chart(scoring_dict, output_path):
    """
    生成雷达图
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import numpy as np

    font_prop = _get_chinese_font_prop()

    # 只画6个子维度
    dims = [k for k in scoring_dict if k != "综合评分"]
    scores = [scoring_dict[k]["score"] for k in dims]

    N = len(dims)
    angles = np.linspace(0, 2 * np.pi, N, endpoint=False).tolist()
    # 闭合
    scores_closed = scores + [scores[0]]
    angles_closed = angles + [angles[0]]

    fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(polar=True))
    ax.fill(angles_closed, scores_closed, color="#7C6AFF", alpha=0.25)
    ax.plot(angles_closed, scores_closed, color="#7C6AFF", linewidth=2)
    ax.set_xticks(angles)
    if font_prop:
        ax.set_xticklabels(dims, fontsize=10, fontproperties=font_prop)
    else:
        ax.set_xticklabels(dims, fontsize=10)
    ax.set_ylim(0, 10)
    ax.set_yticks([2, 4, 6, 8, 10])
    ax.set_yticklabels(["2", "4", "6", "8", "10"], fontsize=8, color="#999999")
    ax.set_title("能力雷达图", fontsize=13, fontweight="bold", pad=20, fontproperties=font_prop)

    # 在各顶点标注分数
    for angle, score, dim in zip(angles, scores, dims):
        ax.annotate(f"{score}", xy=(angle, score), fontsize=10, fontweight="bold",
                     ha="center", va="bottom", color="#333333")

    plt.tight_layout()
    fig.savefig(output_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return output_path


def generate_pdf_report(result, output_path):
    """
    生成带数据可视化图表的 PDF 分析报告
    """
    from fpdf import FPDF

    scoring = get_scoring_data(result)
    ca = result.get("content_analysis", {})

    # 如果有评分数据，先生成图表
    chart_images = {}
    if scoring:
        # 在同目录创建临时图表
        tmp_dir = os.path.dirname(output_path) or "."
        bar_path = os.path.join(tmp_dir, "_tmp_bar.png")
        radar_path = os.path.join(tmp_dir, "_tmp_radar.png")

        try:
            _generate_score_bar_chart(scoring, bar_path)
            chart_images["bar"] = bar_path
        except Exception:
            pass

        try:
            _generate_radar_chart(scoring, radar_path)
            chart_images["radar"] = radar_path
        except Exception:
            pass

    class PDF(FPDF):
        """自定义 PDF 类，支持中文"""
        def __init__(self):
            super().__init__()
            regular_font, bold_font = _get_pdf_font_paths()
            if not regular_font:
                raise RuntimeError("未找到可用于 PDF 导出的 Unicode 字体，请在部署环境安装 Noto Sans CJK 或 DejaVu Sans")
            self.add_font("msyh", "", regular_font, uni=True)
            self.add_font("msyh", "B", bold_font, uni=True)

        def header(self):
            self.set_font("msyh", "B", 10)
            self.set_text_color(160, 160, 160)
            self.cell(0, 8, "信息流素材分析报告", 0, 1, "R")
            self.line(10, 15, 200, 15)
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font("msyh", "", 8)
            self.set_text_color(180, 180, 180)
            self.cell(0, 10, f"第 {self.page_no()}/{{nb}} 页  |  由信息流素材分析 Agent 自动生成", 0, 0, "C")

        def section_title(self, title):
            self.set_font("msyh", "B", 14)
            self.set_text_color(51, 51, 51)
            self.cell(0, 10, title, 0, 1)
            self.set_draw_color(124, 106, 255)
            self.set_line_width(0.6)
            self.line(10, self.get_y(), 60, self.get_y())
            self.ln(4)

        def sub_title(self, title):
            self.set_font("msyh", "B", 11)
            self.set_text_color(80, 80, 80)
            self.cell(0, 8, title, 0, 1)
            self.ln(1)

        def body_text(self, text):
            self.set_x(self.l_margin)
            self.set_font("msyh", "", 10)
            self.set_text_color(60, 60, 60)
            self.multi_cell(0, 6, text)
            self.ln(2)

        def bullet(self, text):
            self.set_x(self.l_margin)
            self.set_font("msyh", "", 10)
            self.set_text_color(60, 60, 60)
            if not text or not text.strip():
                return
            self.multi_cell(0, 6, "  - " + text.strip())

    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)

    # ===== 封面页 =====
    pdf.add_page()
    pdf.ln(50)
    pdf.set_font("msyh", "B", 28)
    pdf.set_text_color(51, 51, 51)
    pdf.cell(0, 15, "信息流素材分析报告", 0, 1, "C")
    pdf.ln(10)

    video_name = result.get("video_name", "未知视频")
    vi = result.get("video_info", {})
    pdf.set_font("msyh", "", 12)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 8, f"视频: {video_name}", 0, 1, "C")
    pdf.cell(0, 8, f"分辨率: {vi.get('width', '?')}x{vi.get('height', '?')}  |  时长: {vi.get('duration', 0)}秒", 0, 1, "C")
    pdf.cell(0, 8, f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 1, "C")

    # 综合评分大字
    if scoring and "综合评分" in scoring:
        overall = scoring["综合评分"]["score"]
        pdf.ln(15)
        pdf.set_font("msyh", "B", 48)
        if overall >= 7:
            pdf.set_text_color(46, 204, 113)
        elif overall >= 4:
            pdf.set_text_color(243, 156, 18)
        else:
            pdf.set_text_color(231, 76, 60)
        pdf.cell(0, 20, f"{overall}", 0, 1, "C")
        pdf.set_font("msyh", "", 12)
        pdf.set_text_color(120, 120, 120)
        pdf.cell(0, 8, "综合评分 (满分10分)", 0, 1, "C")

    # ===== 基本信息页 =====
    pdf.add_page()
    pdf.section_title("一、视频基本信息")
    info_items = [
        ("文件名", video_name),
        ("分辨率", f"{vi.get('width', '?')} x {vi.get('height', '?')}"),
        ("时长", f"{vi.get('duration', 0)} 秒"),
        ("帧率", f"{vi.get('fps', 0)} fps"),
        ("文件大小", f"{vi.get('file_size_mb', 0)} MB"),
        ("音频", "有" if vi.get("has_audio") else "无"),
        ("分析时间", result.get("analyzed_at", "")),
    ]
    pdf.set_font("msyh", "", 10)
    for label, value in info_items:
        pdf.set_font("msyh", "B", 10)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(35, 7, f"{label}:", 0, 0)
        pdf.set_font("msyh", "", 10)
        pdf.set_text_color(60, 60, 60)
        pdf.cell(0, 7, str(value), 0, 1)

    pdf.ln(6)

    # ===== 语音转文字 =====
    pdf.section_title("二、语音转文字结果")
    transcript = result.get("transcript", {})
    if transcript.get("transcribed"):
        pdf.body_text(f"识别模型: {transcript.get('model', '未知')}")
        full_text = transcript.get("full_text", "")
        if full_text:
            pdf.sub_title("识别文本")
            pdf.set_font("msyh", "B", 11)
            pdf.set_text_color(51, 85, 153)
            pdf.multi_cell(0, 7, full_text)
            pdf.ln(3)
    else:
        pdf.body_text("无语音内容或语音识别未执行。")

    # ===== 画面分析 =====
    pdf.section_title("三、关键帧画面分析")
    fa = result.get("frame_analysis", {})
    if fa.get("analyzed"):
        pdf.body_text(f"分析模型: {fa.get('model', '未知')}  |  分析帧数: {fa.get('frames_analyzed', 0)}/{fa.get('total_frames', 0)}")
        if fa.get("overall_visual"):
            pdf.sub_title("整体视觉总结")
            pdf.body_text(fa["overall_visual"])

        frames = fa.get("frames", [])
        if frames:
            pdf.sub_title("逐帧分析")
            for f in frames:
                pdf.set_font("msyh", "B", 10)
                pdf.set_text_color(80, 80, 80)
                pdf.cell(0, 7, f"第 {f.get('frame', '?')} 帧", 0, 1)
                pdf.body_text(f"  内容: {f.get('content', '')}")
                pdf.body_text(f"  文字: {f.get('text', '无')}  |  风格: {f.get('style', '')}  |  变化: {f.get('change', '')}")
    else:
        pdf.body_text("画面分析未执行。")

    # ===== 元素拆解 =====
    if ca.get("analyzed"):
        eb = ca.get("element_breakdown", {})

        pdf.add_page()
        pdf.section_title("四、元素拆解")

        if eb.get("visual"):
            pdf.sub_title("4.1 画面元素")
            vis = eb["visual"]
            for key, label in [("subject", "画面主体"), ("scene", "场景"),
                               ("color_tone", "色调"), ("composition", "构图"), ("effects", "特效")]:
                if vis.get(key):
                    pdf.body_text(f"  {label}: {vis[key]}")

        if eb.get("text"):
            pdf.sub_title("4.2 文案元素")
            txt = eb["text"]
            for key, label in [("spoken", "口播/旁白"), ("subtitles", "字幕/贴纸"),
                               ("cta", "转化引导"), ("headline", "标题/大字")]:
                if txt.get(key):
                    pdf.body_text(f"  {label}: {txt[key]}")

        if eb.get("audio"):
            pdf.sub_title("4.3 音频元素")
            aud = eb["audio"]
            for key, label in [("voice_style", "语音风格"), ("bgm_style", "BGM风格"),
                               ("sound_effects", "音效")]:
                if aud.get(key):
                    pdf.body_text(f"  {label}: {aud[key]}")

        if eb.get("structure"):
            pdf.sub_title("4.4 结构分析")
            st = eb["structure"]
            for key, label in [("hook", "钩子(前3秒)"), ("body", "中间内容"),
                               ("ending", "结尾/CTA"), ("rhythm", "节奏")]:
                if st.get(key):
                    pdf.body_text(f"  {label}: {st[key]}")

        pdf.ln(4)

        # ===== 评分 + 图表 =====
        pdf.section_title("五、评分详情")

        if scoring:
            # 柱状图
            if "bar" in chart_images:
                try:
                    pdf.image(chart_images["bar"], x=15, w=180)
                    pdf.ln(5)
                except Exception:
                    pass

            # 雷达图
            if "radar" in chart_images:
                try:
                    pdf.image(chart_images["radar"], x=55, w=100)
                    pdf.ln(5)
                except Exception:
                    pass

            # 评分明细表
            pdf.sub_title("评分明细")
            pdf.set_font("msyh", "B", 9)
            pdf.set_fill_color(124, 106, 255)
            pdf.set_text_color(255, 255, 255)
            pdf.cell(45, 7, "维度", 1, 0, "C", True)
            pdf.cell(25, 7, "分数", 1, 0, "C", True)
            pdf.cell(120, 7, "评语", 1, 1, "C", True)

            pdf.set_font("msyh", "", 9)
            for key in ["hook_strength", "visual_quality", "copy_quality",
                         "audio_quality", "rhythm", "conversion_guidance", "overall"]:
                dim_label = DIMENSION_LABELS.get(key, key)
                if dim_label in scoring:
                    s = scoring[dim_label]
                    pdf.set_text_color(60, 60, 60)
                    pdf.cell(45, 7, dim_label, 1, 0)
                    pdf.cell(25, 7, f"{s['score']}/{s['max']}", 1, 0, "C")
                    pdf.cell(120, 7, s.get("comment", "")[:40], 1, 1)

        pdf.ln(4)

        # ===== 跑量归因 =====
        pdf.add_page()
        pdf.section_title("六、跑量归因分析")
        ta = ca.get("traffic_attribution", {})

        for section_key, section_label in [
            ("ctr_drivers", "点击率(CTR)驱动因素"),
            ("cvr_drivers", "转化率(CVR)驱动因素"),
            ("completion_drivers", "完播率驱动因素"),
        ]:
            items = ta.get(section_key, [])
            pdf.sub_title(section_label)
            if items:
                for item in items:
                    pdf.bullet(item)
                pdf.ln(2)
            else:
                pdf.body_text("  无明显驱动因素")
                pdf.ln(2)

        issues = ta.get("potential_issues", [])
        if issues:
            pdf.sub_title("潜在问题")
            for item in issues:
                pdf.bullet(item)
            pdf.ln(3)

        if ta.get("overall_assessment"):
            pdf.sub_title("综合跑量评估")
            pdf.body_text(ta["overall_assessment"])

        suggestions = ta.get("improvement_suggestions", [])
        if suggestions:
            pdf.sub_title("改进建议")
            for i, item in enumerate(suggestions, 1):
                pdf.body_text(f"  {i}. {item}")
    else:
        pdf.section_title("四、综合 AI 分析")
        pdf.body_text("AI 综合分析尚未执行，请先完成分析后再导出报告。")

    # 清理临时图表
    for path in chart_images.values():
        try:
            if os.path.exists(path):
                os.remove(path)
        except Exception:
            pass

    # 保存 PDF
    pdf.output(output_path)
    return output_path
